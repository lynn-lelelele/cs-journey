from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_bg(p, color):
    """给段落加背景色"""
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)

def add_h(text, level):
    h = doc.add_heading(text, level=min(level, 4))
    for r in h.runs:
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_p(text, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    set_bg(p, 'F2F2F2')
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_table(headers, rows):
    """headers: list of str, rows: list of list of str"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = '微软雅黑'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = '微软雅黑'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    r.font.size = Pt(10)
    return table

# ======== 正文 ========

add_h('CS50 Lecture 2 · 数组 & 命令行参数', 1)
add_p('2026.08.03 · Lynn')

add_h('1. 命令行参数：argc 和 argv', 2)

add_h('是什么', 3)
add_p('运行程序的时候，除了 ./程序名，还能在后面带参数：')
add_code('./caesar 3')
add_p('这里 3 就是参数。C 语言用 main 的括号来接收它：')
add_code('int main(int argc, string argv[])')

add_h('argc = 参数个数', 3)
add_code('./caesar 3        → argc = 2\n./caesar          → argc = 1\n./caesar 3 hello  → argc = 3')
add_p('规律：argc = 空格分隔的词的数量。程序名本身算一个。')

add_h('argv = 参数数组', 3)
add_code('argv[0] = "./caesar"   ← 永远是程序名\nargv[1] = "3"          ← 第一个参数\nargv[2] = "hello"      ← 第二个参数')
add_p('注意：argv[1] 是字符串，不是数字！ "3" 不等于 3。')

add_h('验证参数数量', 3)
add_code('if (argc != 2)\n{\n    printf("Usage: ./caesar key\\n");\n    return 1;   // 1 表示出错了\n}')

# ----------

add_h('2. 字符串转数字：atoi()', 2)
add_p('argv[1] 是字符串 "3"，不能直接做加减。用 atoi() 转换：')
add_code('#include <stdlib.h>   // atoi 在这里\n\nint key = atoi(argv[1]);   // "3" 变成 3')
add_p('atoi = ASCII to Integer，就是把字符串变成整数。')

# ----------

add_h('3. ASCII 码', 2)
add_p('计算机里所有字符本质都是数字：')
add_code("'A' = 65    'B' = 66    ...    'Z' = 90\n'a' = 97    'b' = 98    ...    'z' = 122\n'0' = 48    '1' = 49    ...    '9' = 57")

add_p('字符可以当数字用：', bold=True)
add_code("char c = 'A';\nint n = c;          // n = 65\n\nchar d = 66;        // d = 'B'\n\nchar e = 'C' + 1;   // e = 'D'")

# ----------

add_h('4. 凯撒加密公式', 2)

add_h('核心思想', 3)
add_p('把 A-Z 这 26 个字母看成一个环。Z 后面回到 A。')
add_code('A B C ... X Y Z A B C ...\n0 1 2     23 24 25')

add_h('四个步骤', 3)
add_p('1. 把字母变成 0-25：s[i] - \'A\'')
add_p('2. 加上密钥：+ key')
add_p('3. 取余 % 26，超了绕回来')
add_p('4. 加回 \'A\'')
add_code('大写：cipher = (s[i] - \'A\' + key) % 26 + \'A\'\n小写：cipher = (s[i] - \'a\' + key) % 26 + \'a\'')

add_h('例子：H + 3', 3)
add_code("'H' = 72\n'H' - 'A' = 72 - 65 = 7\n7 + 3 = 10\n10 % 26 = 10\n10 + 'A' = 65 + 10 = 75 = 'K'")
add_p('H 往后挪 3 位 → K，对了！')

add_h('例子：Z + 1（绕回来）', 3)
add_code("'Z' - 'A' = 25\n25 + 1 = 26\n26 % 26 = 0\n0 + 'A' = 'A'")
add_p('Z 后面回到 A，绕回来了！')

add_h('非字母不管', 3)
add_code('if (isupper(s[i]))        // 大写\n{\n    s[i] = (s[i] - \'A\' + key) % 26 + \'A\';\n}\nelse if (islower(s[i]))   // 小写\n{\n    s[i] = (s[i] - \'a\' + key) % 26 + \'a\';\n}')
add_p('空格、标点、数字——原样保留，不处理。')

add_h('新函数：isupper() 和 islower()', 3)
add_code("isupper('A') → true    isupper('a') → false\nislower('a') → true    islower('A') → false")
add_p('都在 <ctype.h> 里，跟 isalpha() 一个家。')

# ----------

add_h('5. 完整流程', 2)
add_code('./caesar 3\n     ↓\n检查 argc == 2 ？         ← 不是就报错\n     ↓\natoi(argv[1]) 得到 key    ← 把 "3" 变成 3\n     ↓\nget_string 读入明文         ← 昨天学的\n     ↓\n遍历每个字符：\n  - 大写 → 公式加密\n  - 小写 → 公式加密\n  - 其他 → 原样\n     ↓\n打印密文')

# ----------

add_h('6. 常见报错', 2)

add_table(
    ['报错', '原因', '修复'],
    [
        ['segmentation fault', '没检查 argc 就用了 argv[1]', '先写 if (argc != 2)'],
        ['incompatible integer', '把 "3" 直接当数字用了', '用 atoi() 转换'],
        ['implicit declaration of atoi', '没 include stdlib.h', '#include <stdlib.h>'],
        ['输出乱码', '忘了 % 26，值超出字母范围', '公式里加上 % 26'],
    ]
)

# 保存
output = r'C:\Users\19918\Desktop\cs-journey\notes\CS50-Lecture2-数组与命令行参数.docx'
doc.save(output)
print(f'Done! 保存到 {output}')
